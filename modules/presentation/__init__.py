"""AI 专利助手 - 报告导出模块

支持导出为 Markdown 和 Word 文档
"""

import os
import logging
from datetime import datetime
from typing import Optional

from core.output_schema import (
    IdeaMiningResult,
    FiveElements,
    PatentAbstract,
    ClaimSet,
    PatentSpecification,
)

logger = logging.getLogger(__name__)


class ReportExporter:
    """报告导出服务"""

    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def export_full_report(
        self,
        invention_name: str,
        technical_description: str,
        idea_mining_result: IdeaMiningResult,
        five_elements: FiveElements,
        abstract_result: PatentAbstract,
        claims: ClaimSet,
        specification: Optional[PatentSpecification] = None,
    ) -> str:
        """导出完整的专利分析报告为 Markdown

        Returns:
            str: 生成的文件路径
        """
        md_content = self._build_markdown(
            invention_name,
            technical_description,
            idea_mining_result,
            five_elements,
            abstract_result,
            claims,
            specification,
        )

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"专利分析报告_{invention_name[:20]}_{timestamp}.md"
        filepath = os.path.join(self.output_dir, filename)

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(md_content)

        logger.info(f"报告已导出: {filepath}")
        return filepath

    def export_word(
        self,
        invention_name: str,
        technical_description: str,
        idea_mining_result: IdeaMiningResult,
        five_elements: FiveElements,
        abstract_result: PatentAbstract,
        claims: ClaimSet,
        specification: Optional[PatentSpecification] = None,
    ) -> str:
        """导出 Word 文档

        Returns:
            str: 生成的文件路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = Document()

        # 标题
        title = doc.add_heading(invention_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 摘要
        doc.add_heading("摘要", level=1)
        doc.add_paragraph(abstract_result.abstract)
        doc.add_paragraph("关键词：" + "、".join(abstract_result.keywords))

        # 五要素
        doc.add_heading("五要素分析", level=1)
        doc.add_heading("技术问题", level=2)
        doc.add_paragraph(five_elements.technical_problem)
        doc.add_heading("技术方案", level=2)
        doc.add_paragraph(five_elements.technical_solution)
        doc.add_heading("技术效果", level=2)
        doc.add_paragraph(five_elements.technical_effect)
        doc.add_heading("技术特征", level=2)
        for feat in five_elements.technical_features:
            doc.add_paragraph(feat, style="List Bullet")
        doc.add_heading("应用场景", level=2)
        for scene in five_elements.application_scenarios:
            doc.add_paragraph(scene, style="List Bullet")

        # 创新点
        doc.add_heading("创新点分析", level=1)
        for i, inn in enumerate(idea_mining_result.innovations, 1):
            doc.add_heading(f"创新点 {i}: {inn.title}", level=2)
            doc.add_paragraph(f"类型：{inn.innovation_type}")
            doc.add_paragraph(f"描述：{inn.description}")
            doc.add_paragraph(f"技术价值：{inn.technical_value}")
            doc.add_paragraph(f"创新程度：{inn.level}")

        # 新颖性评估
        doc.add_heading("新颖性评估", level=1)
        eval_ = idea_mining_result.evaluation
        doc.add_paragraph(f"新颖性评分：{eval_.novelty_score}/10")
        doc.add_paragraph(f"创造性评分：{eval_.creativity_score}/10")
        doc.add_paragraph(f"市场价值评分：{eval_.market_value_score}/10")
        doc.add_paragraph(f"综合评分：{eval_.overall_score}/10")
        doc.add_paragraph(f"技术进步性：{eval_.technical_progress}")

        # 权利要求书
        doc.add_heading("权利要求书", level=1)
        for claim in claims.claims:
            doc.add_paragraph(f"{claim.claim_number}. {claim.content}")

        # 说明书
        if specification:
            doc.add_heading("专利说明书", level=1)
            doc.add_heading("技术领域", level=2)
            doc.add_paragraph(specification.technical_field)
            doc.add_heading("背景技术", level=2)
            doc.add_paragraph(specification.background_art)
            doc.add_heading("发明内容", level=2)
            doc.add_paragraph(specification.summary)
            doc.add_heading("具体实施方式", level=2)
            doc.add_paragraph(specification.detailed_description)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"专利分析报告_{invention_name[:20]}_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)

        doc.save(filepath)
        logger.info(f"Word 报告已导出: {filepath}")
        return filepath

    def _build_markdown(
        self,
        invention_name: str,
        technical_description: str,
        idea_mining_result: IdeaMiningResult,
        five_elements: FiveElements,
        abstract_result: PatentAbstract,
        claims: ClaimSet,
        specification: Optional[PatentSpecification],
    ) -> str:
        """构建 Markdown 报告内容"""
        sections = []

        sections.append(f"# {invention_name}\n")

        # 摘要
        sections.append("## 摘要\n")
        sections.append(f"{abstract_result.abstract}\n")
        sections.append(f"**关键词**：{'、'.join(abstract_result.keywords)}\n")

        # 五要素
        sections.append("## 五要素分析\n")
        sections.append(f"### 技术问题\n{five_elements.technical_problem}\n")
        sections.append(f"### 技术方案\n{five_elements.technical_solution}\n")
        sections.append(f"### 技术效果\n{five_elements.technical_effect}\n")
        sections.append(f"### 技术特征\n")
        for feat in five_elements.technical_features:
            sections.append(f"- {feat}")
        sections.append(f"### 应用场景\n")
        for scene in five_elements.application_scenarios:
            sections.append(f"- {scene}")
        sections.append("")

        # 创新点
        sections.append("## 创新点分析\n")
        for i, inn in enumerate(idea_mining_result.innovations, 1):
            sections.append(f"### 创新点 {i}: {inn.title}\n")
            sections.append(f"- **类型**：{inn.innovation_type}")
            sections.append(f"- **描述**：{inn.description}")
            sections.append(f"- **技术价值**：{inn.technical_value}")
            sections.append(f"- **创新程度**：{inn.level}")
            sections.append("")

        # 新颖性评估
        eval_ = idea_mining_result.evaluation
        sections.append("## 新颖性评估\n")
        sections.append(f"| 维度 | 评分 |")
        sections.append(f"|------|------|")
        sections.append(f"| 新颖性 | {eval_.novelty_score}/10 |")
        sections.append(f"| 创造性 | {eval_.creativity_score}/10 |")
        sections.append(f"| 市场价值 | {eval_.market_value_score}/10 |")
        sections.append(f"| **综合** | **{eval_.overall_score}/10** |")
        sections.append(f"\n**技术进步性**：{eval_.technical_progress}\n")
        sections.append(f"**优势**：\n")
        for s in eval_.strengths:
            sections.append(f"- {s}")
        sections.append(f"\n**不足/风险**：\n")
        for w in eval_.weaknesses:
            sections.append(f"- {w}")
        sections.append("")

        # 改进建议
        if idea_mining_result.suggestions:
            sections.append("## 改进建议\n")
            for i, sug in enumerate(idea_mining_result.suggestions, 1):
                sections.append(f"### 建议 {i}: {sug.direction}\n")
                sections.append(f"- **建议内容**：{sug.suggestion}")
                sections.append(f"- **预期效果**：{sug.expected_effect}")
                sections.append(f"- **实施难度**：{sug.difficulty}")
                sections.append("")

        # 权利要求书
        sections.append("## 权利要求书\n")
        for claim in claims.claims:
            prefix = f"{claim.claim_number}. "
            sections.append(f"{prefix}{claim.content}\n")

        # 说明书
        if specification:
            sections.append("## 专利说明书\n")
            sections.append(f"### 技术领域\n{specification.technical_field}\n")
            sections.append(f"### 背景技术\n{specification.background_art}\n")
            sections.append(f"### 发明内容\n{specification.summary}\n")
            sections.append(f"### 具体实施方式\n{specification.detailed_description}\n")
            sections.append(f"### 附图说明\n{specification.brief_description_of_drawings}\n")

        return "\n".join(sections)
